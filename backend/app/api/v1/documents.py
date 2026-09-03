"""
Document Library API — Kho tài liệu cá nhân cho từng user
Supports: upload (PDF/DOCX/TXT), list, delete, topic listing
"""
from __future__ import annotations
import os
import uuid
import shutil
from typing import Optional, List
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, distinct

from app.core.dependencies import get_db, get_current_user
from app.models.document import UserDocument, DocumentChunk
from app.schemas.document import DocumentOut, DocumentListItem, TopicTagItem

router = APIRouter(prefix="/documents", tags=["documents"])

# ── Upload directory ──────────────────────────────────────────────────────────
UPLOAD_DIR = Path(__file__).parent.parent.parent.parent / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "md"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB


def _extract_text(file_path: Path, file_type: str) -> str:
    """Extract plain text từ file đã upload"""
    try:
        if file_type == "txt" or file_type == "md":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        elif file_type == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(str(file_path))
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return "\n\n".join(pages)
            except ImportError:
                return "[PDF: Cần cài pypdf để extract text]"
            except Exception as e:
                return f"[Lỗi đọc PDF: {e}]"

        elif file_type == "docx":
            try:
                import docx
                doc = docx.Document(str(file_path))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)
            except ImportError:
                return "[DOCX: Cần cài python-docx để extract text]"
            except Exception as e:
                return f"[Lỗi đọc DOCX: {e}]"

        else:
            return "[Định dạng không hỗ trợ extract text]"
    except Exception as e:
        return f"[Lỗi extract: {e}]"


def _split_into_chunks(text: str, chunk_size: int = 1500) -> List[str]:
    """Chia text thành các chunk ~1500 ký tự để feed vào AI"""
    if not text or not text.strip():
        return []
    words = text.split()
    chunks = []
    current = []
    current_len = 0
    for word in words:
        current.append(word)
        current_len += len(word) + 1
        if current_len >= chunk_size:
            chunks.append(" ".join(current))
            current = []
            current_len = 0
    if current:
        chunks.append(" ".join(current))
    return chunks


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/upload", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(...),
    description: Optional[str] = Form(None),
    topic_tag: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """
    Upload tài liệu (PDF, DOCX, TXT, MD) vào kho cá nhân.
    Server sẽ tự động extract text và chia thành chunks cho AI.
    """
    # Validate extension
    original_name = file.filename or "upload"
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Chỉ chấp nhận file: {', '.join(ALLOWED_EXTENSIONS).upper()}",
        )

    # Read & check size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File quá lớn (tối đa 20MB)")

    # Save to disk
    doc_id = uuid.uuid4()
    user_upload_dir = UPLOAD_DIR / str(current_user.id)
    user_upload_dir.mkdir(parents=True, exist_ok=True)
    save_path = user_upload_dir / f"{doc_id}.{ext}"
    with open(save_path, "wb") as f:
        f.write(content)

    # Extract text
    extracted = _extract_text(save_path, ext)
    chunks = _split_into_chunks(extracted)

    # Persist to DB
    doc = UserDocument(
        id=doc_id,
        user_id=current_user.id,
        title=title.strip(),
        description=description.strip() if description else None,
        file_path=str(save_path),
        original_filename=original_name,
        file_type=ext,
        file_size=len(content),
        topic_tag=topic_tag.strip() if topic_tag else None,
        extracted_text=extracted,
        chunk_count=len(chunks),
        status="active",
    )
    db.add(doc)

    # Persist chunks
    for idx, chunk_text in enumerate(chunks):
        db.add(DocumentChunk(
            document_id=doc_id,
            chunk_index=idx,
            chunk_text=chunk_text,
        ))

    await db.commit()
    await db.refresh(doc)
    return DocumentOut.model_validate(doc)


@router.get("", response_model=List[DocumentListItem])
async def list_documents(
    topic_tag: Optional[str] = None,
    search: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Lấy danh sách tài liệu của user hiện tại (lọc theo topic_tag hoặc search)"""
    stmt = select(UserDocument).where(
        UserDocument.user_id == current_user.id,
        UserDocument.status == "active",
    )
    if topic_tag:
        stmt = stmt.where(UserDocument.topic_tag == topic_tag)
    if search:
        stmt = stmt.where(UserDocument.title.ilike(f"%{search}%"))
    stmt = stmt.order_by(UserDocument.created_at.desc())

    result = await db.execute(stmt)
    docs = result.scalars().all()
    return [DocumentListItem.model_validate(d) for d in docs]


@router.get("/topics", response_model=List[TopicTagItem])
async def list_topics(
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Lấy danh sách các chủ đề (topic_tag) và số tài liệu trong kho của user"""
    stmt = (
        select(UserDocument.topic_tag, func.count(UserDocument.id).label("cnt"))
        .where(
            UserDocument.user_id == current_user.id,
            UserDocument.status == "active",
            UserDocument.topic_tag.isnot(None),
        )
        .group_by(UserDocument.topic_tag)
        .order_by(UserDocument.topic_tag)
    )
    result = await db.execute(stmt)
    rows = result.all()
    return [TopicTagItem(topic_tag=r.topic_tag, document_count=r.cnt) for r in rows]


@router.get("/{doc_id}", response_model=DocumentOut)
async def get_document(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Xem chi tiết tài liệu"""
    stmt = select(UserDocument).where(
        UserDocument.id == doc_id,
        UserDocument.user_id == current_user.id,
        UserDocument.status == "active",
    )
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    return DocumentOut.model_validate(doc)


@router.delete("/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Xóa tài liệu (soft delete + xóa file)"""
    stmt = select(UserDocument).where(
        UserDocument.id == doc_id,
        UserDocument.user_id == current_user.id,
        UserDocument.status == "active",
    )
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")

    # Remove physical file
    if doc.file_path:
        try:
            Path(doc.file_path).unlink(missing_ok=True)
        except Exception:
            pass

    # Soft delete
    doc.status = "deleted"
    await db.commit()
